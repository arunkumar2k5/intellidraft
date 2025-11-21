"""
Document Service - .docx Template Processing
Handles reading template files and updating the Functional Description section
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from typing import Dict, Optional
import shutil
from config import settings

async def update_functional_description(
    template_path: str,
    part_number: str,
    parameters: Dict[str, str],
    component_type: str,
    description: str = ""
) -> dict:
    """
    Update the Functional Description section (section 2) in the .docx template
    
    Args:
        template_path: Path to the template .docx file
        part_number: Component part number (used for output filename)
        parameters: Dictionary of component parameters to insert
        component_type: Type of component
        description: Optional text description
        
    Returns:
        Dictionary with output file path and status
    """
    try:
        # Load the template document
        doc = Document(template_path)
        
        # Simple approach: Add parameter table at the end of the document
        print(f"\n=== Adding parameter table to end of document ===")
        print(f"Part Number: {part_number}")
        print(f"Component Type: {component_type}")
        print(f"Parameters: {len(parameters)} items\n")
        
        # Create parameters table
        # Insert after the description
        table = doc.add_table(rows=1, cols=2)
        
        # Try to apply a table style, fallback to basic grid if not available
        try:
            table.style = 'Light Grid Accent 1'
        except:
            try:
                table.style = 'Table Grid'
            except:
                # If no style works, just use default and add borders manually
                pass
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Parameter'
        header_cells[1].text = 'Value'
        
        # Make header bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                # If no runs exist, create one
                if not paragraph.runs:
                    run = paragraph.add_run(paragraph.text)
                    run.font.bold = True
        
        # Add parameter rows
        for param_name, param_value in parameters.items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(param_name)
            row_cells[1].text = str(param_value)
        
        # Table is already added to the document (no need to move it)
        
        # Generate output filename
        # Clean part number for filename
        safe_part_number = "".join(c for c in part_number if c.isalnum() or c in ('-', '_'))
        output_filename = f"{safe_part_number}.docx"
        output_path = settings.output_dir / output_filename
        
        # Save the document
        doc.save(str(output_path))
        
        return {
            "success": True,
            "output_path": str(output_path),
            "output_filename": output_filename,
            "part_number": part_number,
            "component_type": component_type
        }
        
    except Exception as e:
        print(f"Document generation error: {e}")
        return {
            "error": str(e),
            "success": False
        }

async def validate_template(template_path: str) -> dict:
    """
    Validate that the template file exists and has a Functional Description section
    
    Args:
        template_path: Path to template file
        
    Returns:
        Dictionary with validation status
    """
    try:
        if not Path(template_path).exists():
            return {
                "valid": False,
                "error": "Template file not found"
            }
        
        doc = Document(template_path)
        
        # Check for Functional Description section
        section_found = False
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip().lower()
            if ("2" in text or "2." in text) and "functional description" in text:
                section_found = True
                break
        
        if not section_found:
            return {
                "valid": False,
                "error": "Functional Description section (2.) not found in template"
            }
        
        return {
            "valid": True,
            "message": "Template is valid"
        }
        
    except Exception as e:
        return {
            "valid": False,
            "error": f"Error reading template: {str(e)}"
        }
